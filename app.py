import streamlit as st
import pandas as pd
import os
from docx import Document
from docx.shared import Pt, Inches, Twips
from datetime import datetime
import base64
from io import BytesIO
from lxml import etree

# Set page configuration
st.set_page_config(
    page_title="ระบบคำนวณส่วนแบ่งเงินรางวัลนำจับ",
    page_icon="💰",
    layout="centered"
)

# Add custom CSS
st.markdown("""
<style>
    .main {
        padding: 20px;
    }
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    h1, h2, h3 {
        color: #1E3A8A;
    }
    .stButton>button {
        background-color: #1E3A8A;
        color: white;
        font-weight: bold;
    }
    .stButton>button:hover {
        background-color: #2563EB;
    }
    .info-box {
        background-color: #f8f9fa;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 20px;
    }
    .result-box {
        background-color: #e2f0d9;
        padding: 15px;
        border-radius: 5px;
        margin: 20px 0;
    }
</style>
""", unsafe_allow_html=True)

# Function to load max fine share data
@st.cache_data
def load_max_fine_data():
    # Try different encodings to handle Thai characters
    encodings = ['utf-8-sig', 'utf-8', 'cp874', 'tis-620', 'windows-874']
    
    # Check if file exists
    if not os.path.exists("max_fine_shares.csv"):
        st.error("ไม่พบไฟล์ max_fine_shares.csv กรุณาตรวจสอบว่าไฟล์อยู่ในโฟลเดอร์เดียวกับแอปพลิเคชัน")
        return pd.DataFrame(columns=["พ.ร.บ.", "มาตรา", "จำนวนเงินส่วนแบ่งสูงสุด", "ความผิด"])
    
    for encoding in encodings:
        try:
            df = pd.read_csv("max_fine_shares.csv", encoding=encoding)
            # ตรวจสอบว่าข้อมูลถูกต้องหรือไม่ (มีคอลัมน์ที่ต้องการครบ)
            required_columns = ["พ.ร.บ.", "มาตรา", "จำนวนเงินส่วนแบ่งสูงสุด"]
            
            # ถ้าไม่มีคอลัมน์ที่ต้องการ โปรแกรมอาจอ่านไฟล์ได้แต่ encoding ไม่ถูกต้อง
            if not all(col in df.columns for col in required_columns):
                continue
                
            # แปลงคอลัมน์ 'จำนวนเงินส่วนแบ่งสูงสุด' เป็นตัวเลข
            df['จำนวนเงินส่วนแบ่งสูงสุด'] = pd.to_numeric(df['จำนวนเงินส่วนแบ่งสูงสุด'], errors='coerce')
            
            # เพิ่มคอลัมน์ "ความผิด" หากไม่มี
            if "ความผิด" not in df.columns:
                df["ความผิด"] = ""
                
            # เพิ่มคอลัมน์ "มีจำนวนเงินส่วนแบ่งสูงสุด" เพื่อระบุว่ากฎหมายนี้มีจำนวนเงินส่วนแบ่งสูงสุดหรือไม่
            df['มีจำนวนเงินส่วนแบ่งสูงสุด'] = df['จำนวนเงินส่วนแบ่งสูงสุด'].notna()
                
            return df
            
        except UnicodeDecodeError:
            continue
        except Exception as e:
            pass
    
    st.error("ไม่สามารถอ่านไฟล์ข้อมูลได้ กรุณาตรวจสอบรูปแบบไฟล์และ encoding")
    return pd.DataFrame(columns=["พ.ร.บ.", "มาตรา", "จำนวนเงินส่วนแบ่งสูงสุด", "ความผิด", "มีจำนวนเงินส่วนแบ่งสูงสุด"])

def has_max_share_limit(law_name, section, df):
    """
    ตรวจสอบว่ากฎหมายและมาตราที่ระบุมีจำนวนเงินส่วนแบ่งสูงสุดหรือไม่
    """
    # ถ้าเป็น "มาตรา อื่นๆ" หรือ "ไม่ระบุ" หรือ None ให้ถือว่าไม่มีจำนวนเงินส่วนแบ่งสูงสุด
    if section in ["มาตรา อื่นๆ", "ไม่ระบุ", None]:
        return False, None
    
    # ค้นหากฎหมายและมาตราที่ตรงกัน
    matching_rows = df[(df['พ.ร.บ.'] == law_name) & (df['มาตรา'] == section)]
    
    if matching_rows.empty:
        return False, None
    
    # ตรวจสอบว่ามีจำนวนเงินส่วนแบ่งสูงสุดหรือไม่
    has_limit = matching_rows['มีจำนวนเงินส่วนแบ่งสูงสุด'].iloc[0]
    max_share = matching_rows['จำนวนเงินส่วนแบ่งสูงสุด'].iloc[0] if has_limit else None
    
    return has_limit, max_share

# Function to create and download Word document
def create_word_document(data):
    doc = Document()
    
    # Set page width for the document (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    
    # Add XML parser function
    def parse_xml(xml_string):
        return etree.fromstring(xml_string)
    
    # Set font for the entire document
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(16)
    
    # Add document heading with proper formatting
    title = doc.add_heading("", level=0)
    title_run = title.add_run("ใบสั่งชำระค่าปรับ")
    title_run.font.name = 'TH SarabunPSK'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title.alignment = 1  # Center alignment

    # Add department info (right-aligned)
    header_para1 = doc.add_paragraph()
    header_para1.alignment = 2  # Right alignment
    header_run1 = header_para1.add_run("สำนักงานสาธารณสุขจังหวัดสมุทรปราการ")
    header_run1.font.name = 'TH SarabunPSK'
    header_run1.font.size = Pt(16)

    # Add address (right-aligned)
    header_para2 = doc.add_paragraph()
    header_para2.alignment = 2  # Right alignment
    header_run2 = header_para2.add_run("๑๙ ซอย ๓๕ อัศวนนท์ ๒ สป ๑๐๒๗๐")
    header_run2.font.name = 'TH SarabunPSK'
    header_run2.font.size = Pt(16)

    # Add date field (right-aligned)
    header_para3 = doc.add_paragraph()
    header_para3.alignment = 2  # Right alignment
    header_run3 = header_para3.add_run("วันที่.....................................................")
    header_run3.font.name = 'TH SarabunPSK'
    header_run3.font.size = Pt(16)

    # Add blank space
    doc.add_paragraph()

    # Add recipient line
    recipient_para = doc.add_paragraph()
    recipient_para.add_run("ถึงเจ้าหน้าที่การเงิน กลุ่มงานบริหาร")
    recipient_para.alignment = 0  # Left alignment

    # Add money receipt line
    receipt_from_para = doc.add_paragraph()
    receipt_from_para.add_run("โปรดรับเงินจาก.............................................................................")
    receipt_from_para.alignment = 0  # Left alignment

    
    # Add fine amount
    amount_para = doc.add_paragraph()
    amount_text = f"*จำนวนเงินค่าปรับรวม {data['fine_amount']:,.2f} บาท ({convert_to_thai_text(data['fine_amount'])})"
    amount_para.add_run(amount_text).bold = True
    
    # Add law info
    law_para = doc.add_paragraph()
    law_text = f"เป็นค่าปรับ ตามพระราชบัญญัติ{data['law']} และที่แก้ไขเพิ่มเติม"
    law_para.add_run(law_text)
    
    # Add offense description
    offense_para = doc.add_paragraph()
    if "offense" in data and data["offense"]:
        offense_text = f"ข้อกฎหมายความผิด    {data['offense']} มีบทกำหนดโทษตาม {data['section']}"
    else:
        offense_text = f"ข้อกฎหมายความผิด    .......................................................................................................................................................................................................................................... มีบทกำหนดโทษตามมาตรา {data['section']}"
    offense_para.add_run(offense_text)
    
    # Create a table for the fine calculation (2 columns, 8 rows)
    fine_table = doc.add_table(rows=8, cols=2)
    fine_table.style = 'Table Grid'
    fine_table.autofit = False
    
    # กำหนดความกว้างตาราง 50% ของหน้า
    fine_table._element.tblPr.xpath('./w:tblW')[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '2500')
    fine_table._element.tblPr.xpath('./w:tblW')[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'pct')
    
    # ลบเส้นตารางภายในและใส่เฉพาะเส้นกรอบภายนอก
    tblBorders = parse_xml("""
    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      <w:insideH w:val="nil"/>
      <w:insideV w:val="nil"/>
    </w:tblBorders>
    """)
    
    # ดึง tblPr element
    tblPr = fine_table._element.xpath('./w:tblPr')[0]
    # ลบ tblBorders เดิมถ้ามี
    for element in tblPr.xpath('./w:tblBorders'):
        tblPr.remove(element)
    # เพิ่ม tblBorders ใหม่
    tblPr.append(tblBorders)
    
    # Add the box title
    fine_box_cell = fine_table.cell(0, 0)
    fine_box_cell.merge(fine_table.cell(0, 1))
    fine_box_para = fine_box_cell.paragraphs[0]
    fine_box_para.add_run(f"กันเงิน...60...%*").bold = True
    fine_box_para.alignment = 1  # Center
    
    # Add calculation rows
    fine_table.cell(1, 0).text = "จำนวนเงิน"
    fine_table.cell(1, 1).text = f"{data['calculated_share']:,.2f} บาท"
    
    fine_table.cell(2, 0).text = "สูงสุดไม่เกิน"
    if data['max_share'] == float('inf'):
        fine_table.cell(2, 1).text = "ไม่มีกำหนด"
    else:
        fine_table.cell(2, 1).text = f"{data['max_share']:,.2f} บาท"
    
    fine_table.cell(3, 0).text = "เงินสินบนนำจับ"
    fine_table.cell(3, 1).text = f"{data['share1']:,.2f} บาท(25 %*)"
    
    # Add checkboxes based on bounty claimant status and law type
    check_cell = fine_table.cell(4, 0)
    check_cell.merge(fine_table.cell(4, 1))
    check_para = check_cell.paragraphs[0]
    
    if data.get('has_bounty_claimant', False):
        check_para.add_run("☑ จ่ายให้ผู้ขอรับสินบนนำจับ")
    else:
        check_para.add_run("□ จ่ายให้ผู้ขอรับสินบนนำจับ")
    
    check_cell2 = fine_table.cell(5, 0)
    check_cell2.merge(fine_table.cell(5, 1))
    check_para2 = check_cell2.paragraphs[0]
    
    if not data.get('has_bounty_claimant', False):
        if data['law'] in ['เครื่องสำอาง', 'เครื่องมือแพทย์']:
            check_para2.add_run("☑ เป็นรายได้แผ่นดิน")
        else:
            check_para2.add_run("☑ รวมกับสินบนรางวัล")
    else:
        check_para2.add_run("□ เป็นรายได้แผ่นดิน")
    
    # Add reward and expense rows
    fine_table.cell(6, 0).text = "รางวัล"
    fine_table.cell(6, 1).text = f"{data['share2']:,.2f} บาท(50 %*)"
    
    fine_table.cell(7, 0).text = "คชจ."
    fine_table.cell(7, 1).text = f"{data['share3']:,.2f} บาท(25 %*)"
    
    # Set font for all cells
    for row in fine_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'TH SarabunPSK'
                paragraph.alignment = 1  # Center align
    
    # Add signature section
    doc.add_paragraph()
    sig_section = doc.add_paragraph()
    sig_section.alignment = 2  # Right alignment
    sig_section.add_run("ผู้รับชำระ.........................................\n")
    sig_section.add_run("โทร ................................................")
    
    # Save to BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Function to convert number to Thai text
def convert_to_thai_text(number):
    # A simple implementation to convert numbers to Thai text
    # This is a basic implementation and might need more sophistication for real use
    
    if number == 0:
        return "ศูนย์บาทถ้วน"
    
    # Split into integer and decimal parts
    integer_part = int(number)
    decimal_part = int(round((number - integer_part) * 100))
    
    # Thai digits
    thai_digits = ["", "หนึ่ง", "สอง", "สาม", "สี่", "ห้า", "หก", "เจ็ด", "แปด", "เก้า"]
    
    # Thai units
    thai_units = ["", "สิบ", "ร้อย", "พัน", "หมื่น", "แสน", "ล้าน"]
    
    # Convert integer part
    result = ""
    
    if integer_part >= 1000000:
        millions = integer_part // 1000000
        result += convert_to_thai_text(millions) + "ล้าน"
        integer_part %= 1000000
    
    # Process each digit
    digits = [int(d) for d in str(integer_part)]
    length = len(digits)
    
    for i in range(length):
        digit = digits[i]
        if digit == 0:
            continue
            
        if i == length - 1 and digit == 1 and length > 1:
            result += "เอ็ด"
        elif i == length - 2 and digit == 2:
            result += "ยี่สิบ"
        elif i == length - 2 and digit == 1:
            result += "สิบ"
        else:
            result += thai_digits[digit] + thai_units[length - i - 1]
    
    # Add "baht"
    result += "บาท"
    
    # Add decimal part if exists
    if decimal_part > 0:
        if decimal_part < 10:
            result += thai_digits[decimal_part] + "สตางค์"
        else:
            tens = decimal_part // 10
            ones = decimal_part % 10
            
            if tens == 2:
                result += "ยี่สิบ"
            elif tens == 1:
                result += "สิบ"
            else:
                result += thai_digits[tens] + "สิบ"
                
            if ones == 1:
                result += "เอ็ดสตางค์"
            elif ones > 0:
                result += thai_digits[ones] + "สตางค์"
            else:
                result += "สตางค์"
    else:
        result += "ถ้วน"
        
    return result

# Function to get download link for docx
def get_download_link(buffer, filename="รายงานการคำนวณส่วนแบ่งเงินรางวัลนำจับ.docx"):
    b64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 ดาวน์โหลดรายงาน Word</a>'

# Main function
def main():
    st.title("💰 ระบบคำนวณส่วนแบ่งเงินรางวัลนำจับ")
    
    # Load max fine data
    df = load_max_fine_data()
    
    # Get unique laws from the data
    laws = ["กรุณาเลือก..."] + df["พ.ร.บ."].unique().tolist()
    
    with st.container():
        st.markdown('<div class="info-box">', unsafe_allow_html=True)
        st.subheader("📝 กรอกข้อมูลเพื่อคำนวณส่วนแบ่ง")
        
        # Input for fine amount
        fine_amount = st.number_input("จำนวนเงินค่าปรับ (บาท)", min_value=0, value=None, step=100, placeholder="กรอกจำนวนเงิน")
        
        # Select law
        selected_law = st.selectbox("เลือกพระราชบัญญัติ", laws)
        
        # Filter sections based on selected law
        if selected_law == "กรุณาเลือก...":
            sections = ["กรุณาเลือก..."]
            filtered_sections = pd.DataFrame(columns=df.columns)
        else:
            filtered_sections = df[df["พ.ร.บ."] == selected_law]
            # Replace NaN values with "ไม่ระบุ" for sections
            sections = ["กรุณาเลือก..."] + [section if pd.notna(section) else "ไม่ระบุ" for section in filtered_sections["มาตรา"].tolist()]
        
        # Select section
        selected_section = st.selectbox("เลือกบทกำหนดโทษ", sections)
        
        # Add checkbox for bounty claimant
        has_bounty_claimant = st.checkbox("มีผู้ขอรับสินบนนำจับ")
        
        # Get offense information if available
        offense_info = ""
        if selected_section != "กรุณาเลือก..." and selected_law != "กรุณาเลือก...":
            # Handle the case where section is "ไม่ระบุ" or "มาตรา อื่นๆ"
            section_to_match = None if selected_section in ["ไม่ระบุ", "มาตรา อื่นๆ"] else selected_section
            selected_row = filtered_sections[filtered_sections["มาตรา"] == section_to_match]
            if not selected_row.empty and "ความผิด" in selected_row.columns:
                offense_info = selected_row["ความผิด"].values[0]
                if pd.notna(offense_info) and offense_info:
                    st.info(f"**ความผิด**: {offense_info}")

        st.markdown('</div>', unsafe_allow_html=True)
        
        # Calculate button
        if st.button("คำนวณส่วนแบ่ง"):
            if fine_amount is None or fine_amount <= 0:
                st.error("กรุณากรอกจำนวนเงินค่าปรับมากกว่า 0 บาท")
            elif selected_law == "กรุณาเลือก...":
                st.error("กรุณาเลือกพระราชบัญญัติ")
            elif selected_section == "กรุณาเลือก...":
                st.error("กรุณาเลือกบทกำหนดโทษ")
            else:
                # Calculate 60% of fine
                calculated_share = fine_amount * 0.6
                
                # Get maximum share for selected law and section
                # Handle the case where section is "ไม่ระบุ" or "มาตรา อื่นๆ"
                section_to_match = None if selected_section in ["ไม่ระบุ", "มาตรา อื่นๆ"] else selected_section
                max_share_row = filtered_sections[filtered_sections["มาตรา"] == section_to_match]
                
                # Check if the law has a maximum share limit
                has_limit, max_share = has_max_share_limit(selected_law, section_to_match, df)
                
                if not has_limit:
                    # If no maximum share limit, use the calculated share directly
                    actual_share = calculated_share
                    max_share_display = "ไม่มีกำหนด"
                else:
                    # Use the minimum between calculated share and maximum share
                    actual_share = min(calculated_share, max_share)
                    max_share_display = f"{max_share:,.2f} บาท"
                
                # Calculate distribution
                share1 = actual_share * 0.25  # 25% - เงินสินบนนำจับ
                share2 = actual_share * 0.50  # 50% - เงินรางวัล
                share3 = actual_share * 0.25  # 25% - ค่าใช้จ่ายในการดำเนินงาน
                
                # Display results
                st.markdown('<div class="result-box">', unsafe_allow_html=True)
                st.subheader("💵 ผลการคำนวณ")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"จำนวนเงินค่าปรับ: **{fine_amount:,.2f}** บาท")
                    st.write(f"ส่วนแบ่งที่คำนวณได้ (60%): **{calculated_share:,.2f}** บาท")
                    st.write(f"จำนวนเงินส่วนแบ่งสูงสุดตามบทกำหนดโทษ: **{max_share_display}**")
                
                with col2:
                    st.write(f"จำนวนเงินส่วนแบ่งที่ใช้จริง: **{actual_share:,.2f}** บาท")
                    st.write(f"พระราชบัญญัติ: **{selected_law}**")
                    st.write(f"บทกำหนดโทษ: **{selected_section}**")
                
                st.markdown("---")
                st.subheader("🔄 การแบ่งส่วนรางวัลนำจับ")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("ส่วนที่ 1 (25%) เงินสินบนนำจับ", f"{share1:,.2f} บาท")
                
                with col2:
                    st.metric("ส่วนที่ 2 (50%) เงินรางวัล", f"{share2:,.2f} บาท")
                
                with col3:
                    st.metric("ส่วนที่ 3 (25%) ค่าใช้จ่ายในการดำเนินงาน", f"{share3:,.2f} บาท")
                
                st.markdown("---")
                
                # Create Word document
                data = {
                    "law": selected_law,
                    "section": "..............................................................................................." if selected_section in ["ไม่ระบุ", "มาตรา อื่นๆ"] else selected_section,
                    "fine_amount": fine_amount,
                    "max_share": max_share if has_limit else float('inf'),
                    "calculated_share": calculated_share,
                    "actual_share": actual_share,
                    "share1": share1,
                    "share2": share2,
                    "share3": share3,
                    "has_bounty_claimant": has_bounty_claimant
                }
                
                # เพิ่มข้อมูลความผิด (ถ้ามี)
                if "ความผิด" in max_share_row.columns and not max_share_row.empty:
                    offense_text = max_share_row["ความผิด"].values[0]
                    if pd.notna(offense_text) and offense_text:
                        data["offense"] = offense_text
                    else:
                        data["offense"] = "..............................................................................................."
                else:
                    data["offense"] = "..............................................................................................."
                
                buffer = create_word_document(data)
                
                # Provide download link
                st.markdown(get_download_link(buffer), unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main() 